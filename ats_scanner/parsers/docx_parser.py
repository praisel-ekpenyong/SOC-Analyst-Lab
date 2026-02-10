"""DOCX parsing with python-docx."""

import re
from typing import List, Tuple
from docx import Document
from ..models.resume import ParsingWarning, SeverityLevel


class DOCXParser:
    """DOCX parser using python-docx."""
    
    def __init__(self):
        self.warnings: List[ParsingWarning] = []
    
    def parse(self, file_path: str) -> Tuple[str, List[ParsingWarning]]:
        """Parse DOCX file and return text and warnings."""
        self.warnings = []
        
        try:
            doc = Document(file_path)
            paragraphs = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            
            full_text = "\n".join(paragraphs)
            
        except Exception as e:
            self.warnings.append(ParsingWarning(
                message=f"Error parsing DOCX: {str(e)}",
                severity=SeverityLevel.HIGH,
                example=None
            ))
            return "", self.warnings
        
        # Run quality checks
        self._check_parsing_quality(full_text)
        
        # Normalize text
        normalized_text = self._normalize_text(full_text)
        
        return normalized_text, self.warnings
    
    def _check_parsing_quality(self, full_text: str):
        """Check for parsing quality issues."""
        # Check for low text extraction
        if len(full_text.strip()) < 100:
            self.warnings.append(ParsingWarning(
                message="Very little text extracted from DOCX",
                severity=SeverityLevel.HIGH,
                example=full_text[:50] if full_text else None
            ))
        
        # Check for excessive special characters
        special_chars = re.findall(r'[^\w\s\-\.,;:()@/]', full_text)
        if len(special_chars) > len(full_text) * 0.1:
            self.warnings.append(ParsingWarning(
                message="Excessive special characters detected",
                severity=SeverityLevel.MEDIUM,
                example=f"Found {len(special_chars)} special characters"
            ))
        
        # Check for missing common section headers
        if not self._has_common_sections(full_text):
            self.warnings.append(ParsingWarning(
                message="Missing common resume section headers",
                severity=SeverityLevel.MEDIUM,
                example="Expected sections like Experience, Education, Skills not found"
            ))
    
    def _has_common_sections(self, text: str) -> bool:
        """Check if text contains common resume sections."""
        text_lower = text.lower()
        common_sections = [
            'experience', 'education', 'skills', 'work history',
            'employment', 'professional', 'summary', 'objective'
        ]
        return any(section in text_lower for section in common_sections)
    
    def _normalize_text(self, text: str) -> str:
        """Normalize whitespace and fix line wraps."""
        # Fix hyphenated line breaks
        text = re.sub(r'(\w)-\s*\n\s*(\w)', r'\1\2', text)
        
        # Normalize whitespace
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        return text.strip()
